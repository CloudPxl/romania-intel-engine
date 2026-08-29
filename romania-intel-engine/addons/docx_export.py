"""Shared DOCX rendering utility for the drafting engines.

The dossier/FOIA/clarification generators produce structured Python data
(headers, ordered sections, compliance rows) — this module is the only
place that turns that structure into a formatted .docx file, so every
exported document shares the same cover page, margins, heading style,
table formatting and signature block instead of each caller reinventing
its own layout.

Romanian public-tender submissions are conventionally bound/printed with
2 cm side margins and 2.5 cm top/bottom (the convention most authorities'
own model documentation and ANAP templates use) — set once here rather
than per document.
"""
import io
import logging
from datetime import date
from typing import Any, Dict, List, Optional, Sequence

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger("DocxExport")

BRAND_COLOR = RGBColor(0x24, 0x2A, 0x88)  # matches the frontend's brand-900 indigo
MARGIN_TOP_BOTTOM_CM = 2.5
MARGIN_SIDES_CM = 2.0


def _set_margins(document: Document) -> None:
    for section in document.sections:
        section.top_margin = Cm(MARGIN_TOP_BOTTOM_CM)
        section.bottom_margin = Cm(MARGIN_TOP_BOTTOM_CM)
        section.left_margin = Cm(MARGIN_SIDES_CM)
        section.right_margin = Cm(MARGIN_SIDES_CM)


def _add_field(paragraph, field_code: str) -> None:
    """Inserts a Word field code (used for TOC and PAGE fields) — python-docx
    has no high-level API for these, so the raw OOXML run is built directly."""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(fld_end)


def _style_document(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for level, size in ((1, 15), (2, 13)):
        style = document.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BRAND_COLOR


def _add_cover_page(
    document: Document,
    *,
    doc_title: str,
    subtitle: str,
    company_name: str,
    cui: Optional[str],
    authority_name: str,
    reference_id: Optional[str],
    generated_on: Optional[date] = None,
) -> None:
    generated_on = generated_on or date.today()

    for _ in range(4):
        document.add_paragraph()

    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(doc_title.upper())
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = BRAND_COLOR

    subtitle_p = document.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_p.add_run(subtitle)
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.italic = True

    for _ in range(6):
        document.add_paragraph()

    meta_rows = [
        ("Autoritate contractantă", authority_name),
        ("Ofertant / Solicitant", company_name),
    ]
    if cui:
        meta_rows.append(("CUI", cui))
    if reference_id:
        meta_rows.append(("Referință procedură", reference_id))
    meta_rows.append(("Data întocmirii", generated_on.strftime("%d.%m.%Y")))

    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for label, value in meta_rows:
        row = table.add_row()
        row.cells[0].width = Cm(6)
        row.cells[1].width = Cm(9)
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(str(value))

    for _ in range(4):
        document.add_paragraph()

    footer_p = document.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run(
        "Generat de RO-INTEL — platformă de inteligență în achiziții publice"
    )
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    document.add_page_break()


def _add_table_of_contents(document: Document, section_titles: Sequence[str]) -> None:
    heading = document.add_heading("CUPRINS", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    toc_p = document.add_paragraph()
    _add_field(
        toc_p,
        r'TOC \o "1-2" \h \z \u',
    )
    # Word only regenerates the TOC field on open (F9 / "Update Field"),
    # so a plain numbered listing is written alongside it as a fallback
    # that's readable immediately in any viewer, including ones that
    # don't evaluate fields at all.
    for i, title in enumerate(section_titles, start=1):
        line = document.add_paragraph(f"{i}. {title}")
        line.paragraph_format.left_indent = Cm(0.5)

    document.add_page_break()


def _add_compliance_table(document: Document, rows: List[Dict[str, str]]) -> None:
    """rows: list of {"requirement": ..., "response": ..., "reference": ...}"""
    table = document.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Cerință din documentația de atribuire"
    hdr[1].text = "Modul de îndeplinire / Răspuns ofertant"
    hdr[2].text = "Referință / Pagină document justificativ"

    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    for row in rows:
        cells = table.add_row().cells
        cells[0].text = str(row.get("requirement", ""))
        cells[1].text = str(row.get("response", ""))
        cells[2].text = str(row.get("reference", ""))


def _add_signature_block(document: Document, *, company_name: str, signatory_role: str = "Reprezentant legal") -> None:
    document.add_paragraph()
    document.add_paragraph()
    p = document.add_paragraph()
    p.add_run(company_name).bold = True
    document.add_paragraph(signatory_role)
    document.add_paragraph()
    sig_p = document.add_paragraph()
    sig_p.add_run("Semnătura: ____________________________")
    date_p = document.add_paragraph()
    date_p.add_run(f"Data: {date.today().strftime('%d.%m.%Y')}")
    document.add_paragraph()
    stamp_p = document.add_paragraph()
    stamp_p.add_run("L.S. (loc pentru ștampilă, dacă este cazul)")
    stamp_p.runs[0].font.italic = True
    stamp_p.runs[0].font.size = Pt(9)


def _write_body_sections(document: Document, sections: Sequence[Dict[str, Any]]) -> None:
    """sections: list of {"heading": str, "paragraphs": [str, ...]}"""
    for section in sections:
        document.add_heading(section["heading"], level=1)
        for para_text in section.get("paragraphs", []):
            if not para_text:
                continue
            document.add_paragraph(para_text)


def render_docx(
    *,
    doc_title: str,
    subtitle: str,
    company_name: str,
    authority_name: str,
    sections: Sequence[Dict[str, Any]],
    cui: Optional[str] = None,
    reference_id: Optional[str] = None,
    compliance_rows: Optional[List[Dict[str, str]]] = None,
    disclaimer: Optional[str] = None,
    signatory_role: str = "Reprezentant legal",
) -> bytes:
    """Assembles one formatted .docx and returns it as raw bytes, ready to
    stream back over HTTP. Shared by both the technical-dossier export and
    the clarification/FOIA letter export so both get the same cover page,
    margins, ToC and signature-block conventions."""
    document = Document()
    _set_margins(document)
    _style_document(document)

    _add_cover_page(
        document,
        doc_title=doc_title,
        subtitle=subtitle,
        company_name=company_name,
        cui=cui,
        authority_name=authority_name,
        reference_id=reference_id,
    )

    section_titles = [s["heading"] for s in sections]
    if compliance_rows:
        section_titles.append("Matricea de conformitate")
    section_titles.append("Semnături")
    _add_table_of_contents(document, section_titles)

    _write_body_sections(document, sections)

    if compliance_rows:
        document.add_heading("Matricea de conformitate", level=1)
        _add_compliance_table(document, compliance_rows)

    if disclaimer:
        document.add_paragraph()
        disc_p = document.add_paragraph()
        disc_run = disc_p.add_run(disclaimer)
        disc_run.font.italic = True
        disc_run.font.size = Pt(9)
        disc_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    document.add_heading("Semnături", level=1)
    _add_signature_block(document, company_name=company_name, signatory_role=signatory_role)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
